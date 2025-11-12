import os
import io
import tempfile
from datetime import datetime
from urllib.parse import quote

import requests
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from openai import OpenAI
from supabase import create_client, Client
import pypandoc

# ============================================================
# CONFIGURAÇÃO INICIAL
# ============================================================

# Config da página (apenas uma vez no app inteiro)
st.set_page_config(
    page_title="Ferramenta Inteligente para Elaboração de ETP",
    layout="wide"
)

# Lê variáveis de ambiente / secrets do Streamlit
SUPABASE_URL = os.getenv("SUPABASE_URL") or st.secrets.get("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY") or st.secrets.get("SUPABASE_KEY")
APP_BASE_URL = (
    os.getenv("APP_BASE_URL")
    or st.secrets.get("APP_BASE_URL")
    or "https://etp-com-ia.streamlit.app"
)

if not SUPABASE_URL or not SUPABASE_KEY:
    st.error("SUPABASE_URL e/ou SUPABASE_KEY não configuradas.")
    st.stop()

# Cliente Supabase
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ============================================================
# FIX: OAuth Google em iframe (#access_token -> ?access_token)
# ============================================================

# Isso resolve o problema do Google devolver o token no hash (#)
# quando o app está em iframe (Streamlit Cloud). Trocamos por query (?).
components.html(
    """
    <script>
    (function () {
      try {
        var w = window.parent || window.top || window;
        var h = w.location.hash || "";
        if (h && h.indexOf("access_token=") >= 0) {
          var qs = h.substring(1);
          var base = w.location.origin + w.location.pathname;
          w.history.replaceState({}, "", base + "?" + qs);
          w.location.reload();
        }
      } catch (e) {
        console.warn("hash->query (parent) error", e);
      }
    })();
    </script>
    """,
    height=0,
)

# ============================================================
# CONSTANTES: ETAPAS / ORIENTAÇÕES / CAMPOS
# ============================================================

ETAPAS = [
    (1, "Ajuste da Descrição da Necessidade de Contratação"),
    (2, "Requisitos da Contratação"),
    (3, "Levantamento de Mercado"),
    (4, "Descrição da solução como um todo"),
    (5, "Estimativa das quantidades"),
    (6, "Estimativa do valor da contratação"),
    (7, "Alinhamento da contratação com PCA"),
    (8, "Justificativa para o parcelamento ou não da contratação"),
    (9, "Contratações correlatas e/ou interdependentes"),
    (10, "Gestão de riscos / riscos envolvidos"),
    (11, "Justificativa da escolha da solução"),
    (12, "Providências finais / conclusão"),
]

ORIENTACOES = {
    1: "Explique o problema, a demanda e o contexto que justificam a contratação.",
    2: "Liste os requisitos funcionais e não funcionais, requisitos legais e restrições.",
    3: "Descreva as pesquisas de mercado, fornecedores consultados, tecnologias existentes.",
    4: "Apresente a solução como um todo, de forma clara e compreensível para não técnicos.",
    5: "Estime as quantidades envolvidas (unidades, horas, licenças, etc.).",
    6: "Detalhe a metodologia de estimativa de valor (cotações, bancos de preços, etc.).",
    7: "Mostre como a contratação está alinhada ao PCA / planejamento institucional.",
    8: "Justifique o parcelamento ou a contratação em lote único, com base na legislação.",
    9: "Indique contratações relacionadas, dependências e impactos interdependentes.",
    10: "Identifique os riscos e as medidas de mitigação associados à contratação.",
    11: "Justifique a escolha da solução em relação a alternativas e critérios adotados.",
    12: "Faça o resumo final e consolide as principais conclusões do ETP.",
}

INFOS_BASICAS_CAMPOS = [
    ("orgao", "Órgão / Entidade"),
    ("unidade", "Unidade Demandante"),
    ("processo", "Número do Processo"),
    ("responsavel", "Responsável pela Demanda"),
    ("objeto", "Objeto da Contratação (resumo)"),
]

# ============================================================
# HELPERS DE AUTENTICAÇÃO (Google + Email/Senha)
# ============================================================


def gerar_google_auth_url() -> str:
    """
    Monta a URL de autenticação do Google via Supabase Auth.
    """
    redirect_enc = quote(APP_BASE_URL, safe="")
    return (
        f"{SUPABASE_URL}/auth/v1/authorize"
        f"?provider=google&redirect_to={redirect_enc}"
    )


def obter_user_supabase(access_token: str):
    """
    Usa o token JWT (access_token) para consultar /auth/v1/user no Supabase
    e obter os dados do usuário.
    """
    if not access_token:
        return None
    try:
        headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {access_token}",
        }
        r = requests.get(
            f"{SUPABASE_URL}/auth/v1/user", headers=headers, timeout=10
        )
        if r.status_code == 200:
            return r.json()
    except Exception:
        return None
    return None


def obter_usuario_por_email(email: str):
    """
    Consulta a tabela 'usuarios' do seu banco, para sincronizar dados
    de login com seu modelo de usuários.
    """
    if not email:
        return None
    resp = (
        supabase.table("usuarios")
        .select("*")
        .eq("email", email)
        .limit(1)
        .execute()
    )
    data = resp.data or []
    return data[0] if data else None


def criar_usuario(nome: str, sobrenome: str, cpf: str, email: str):
    """
    Cria um registro na tabela 'usuarios' do seu banco de dados.
    """
    return (
        supabase.table("usuarios")
        .insert(
            {
                "nome": nome,
                "sobrenome": sobrenome,
                "cpf": cpf,
                "email": email,
            }
        )
        .execute()
        .data[0]
    )


def sincronizar_usuario(user_json: dict):
    """
    Recebe o user_json do Supabase Auth e garante que exista um
    registro correspondente na tabela 'usuarios'.
    """
    if not user_json:
        return None

    email = user_json.get("email")
    meta = user_json.get("user_metadata") or {}
    nome_completo = meta.get("full_name") or meta.get("name") or ""

    partes = nome_completo.split(" ", 1)
    nome = partes[0] if partes else ""
    sobrenome = partes[1] if len(partes) > 1 else ""

    existente = obter_usuario_por_email(email) if email else None
    return existente or criar_usuario(nome, sobrenome, "", email)


# ============================================================
# DB HELPERS (RLS-friendly, sempre filtra por user_id)
# ============================================================


def _require_session():
    """
    Garante que existe um access_token em sessão,
    aplica no postgrest (RLS) e retorna o token.
    """
    token = st.session_state.get("access_token")
    if not token:
        st.warning("Sessão expirada. Faça login novamente.")
        st.session_state.clear()
        st.experimental_set_query_params()
        st.rerun()
    # Importante para RLS: todas as queries abaixo rodam
    # com o JWT do usuário autenticado:
    supabase.postgrest.auth(token)
    return token


def listar_projetos():
    """
    Lista apenas os projetos do usuário logado.
    """
    _require_session()
    user_id = st.session_state.get("auth_user_id")
    if not user_id:
        return []
    return (
        supabase.table("projetos")
        .select("id, nome, criado_em")
        .eq("user_id", user_id)
        .order("criado_em", desc=True)
        .execute()
    ).data or []


def criar_projeto(nome: str):
    """
    Cria um novo projeto atrelado ao usuário logado.
    """
    _require_session()
    user_id = st.session_state.get("auth_user_id")
    if not user_id:
        st.error("Usuário não autenticado.")
        st.stop()
    return (
        supabase.table("projetos")
        .insert({"nome": nome, "user_id": user_id})
        .execute()
        .data[0]["id"]
    )


def obter_projeto(projeto_id: int):
    """
    Carrega os dados básicos de um projeto específico.
    (A RLS no Supabase deve garantir que o usuário só
    veja projetos do próprio user_id.)
    """
    _require_session()
    return (
        supabase.table("projetos")
        .select("*")
        .eq("id", projeto_id)
        .single()
        .execute()
    ).data


def excluir_projeto(projeto_id: int):
    """
    Exclui um projeto específico, garantindo que seja
    do usuário logado (user_id).
    """
    _require_session()
    user_id = st.session_state.get("auth_user_id")
    (
        supabase.table("projetos")
        .delete()
        .eq("id", projeto_id)
        .eq("user_id", user_id)
        .execute()
    )


def atualizar_infos_basicas(projeto_id: int, dados: dict):
    """
    Atualiza as informações básicas (somente daquele projeto).
    IMPORTANTE: filtramos por id e user_id para não afetar
    outros projetos do mesmo usuário.
    """
    _require_session()
    user_id = st.session_state.get("auth_user_id")
    (
        supabase.table("projetos")
        .update(
            {
                "orgao": dados.get("orgao"),
                "unidade": dados.get("unidade"),
                "processo": dados.get("processo"),
                "responsavel": dados.get("responsavel"),
                "objeto": dados.get("objeto"),
            }
        )
        .eq("id", projeto_id)
        .eq("user_id", user_id)
        .execute()
    )


def carregar_etapa(projeto_id: int, numero: int):
    """
    Carrega o texto da etapa específica (texto_final, sugestao_ia, título).
    """
    _require_session()
    resp = (
        supabase.table("etapas")
        .select("texto_final, sugestao_ia, titulo")
        .eq("projeto_id", projeto_id)
        .eq("numero", numero)
        .execute()
    ).data
    if resp:
        row = resp[0]
        return {
            "texto_final": row.get("texto_final") or "",
            "sugestao_ia": row.get("sugestao_ia") or "",
            "titulo": row.get("titulo") or dict(ETAPAS)[numero],
        }
    return {"texto_final": "", "sugestao_ia": "", "titulo": dict(ETAPAS)[numero]}


def salvar_etapa(
    projeto_id: int,
    numero: int,
    titulo: str,
    texto_final: str,
    sugestao_ia: str,
):
    """
    Upsert da etapa (se existe, atualiza; senão, insere).
    """
    _require_session()
    payload = {
        "projeto_id": projeto_id,
        "numero": numero,
        "titulo": titulo,
        "texto_final": texto_final,
        "sugestao_ia": sugestao_ia,
        "atualizado_em": datetime.utcnow().isoformat(),
    }
    existe = (
        supabase.table("etapas")
        .select("id")
        .eq("projeto_id", projeto_id)
        .eq("numero", numero)
        .execute()
    ).data
    if existe:
        (
            supabase.table("etapas")
            .update(payload)
            .eq("projeto_id", projeto_id)
            .eq("numero", numero)
            .execute()
        )
    else:
        supabase.table("etapas").insert(payload).execute()


def salvar_arquivo(projeto_id: int, numero_etapa: int, file):
    """
    Registra metadados do arquivo enviado (nome, etapa etc.).
    """
    _require_session()
    supabase.table("arquivos").insert(
        {
            "projeto_id": projeto_id,
            "numero_etapa": numero_etapa,
            "nome_original": file.name,
            "upload_em": datetime.utcnow().isoformat(),
        }
    ).execute()


def listar_arquivos(projeto_id: int, numero_etapa: int):
    """
    Lista arquivos associados à etapa de um projeto.
    """
    _require_session()
    return (
        supabase.table("arquivos")
        .select("id, nome_original")
        .eq("projeto_id", projeto_id)
        .eq("numero_etapa", numero_etapa)
        .order("upload_em", desc=True)
        .execute()
    ).data or []


def carregar_textos_todas_etapas(projeto_id: int):
    """
    Carrega o texto_final de todas as etapas de um projeto,
    para exportar DOCX/PDF.
    """
    _require_session()
    return (
        supabase.table("etapas")
        .select("numero, titulo, texto_final")
        .eq("projeto_id", projeto_id)
        .order("numero")
        .execute()
    ).data or []

# ============================================================
# IA (OpenAI Responses API)
# ============================================================


def gerar_texto_ia(
    numero_etapa: int,
    nome_etapa: str,
    orientacao: str,
    texto_existente: str,
    infos_basicas: dict,
    arquivos_etapa: list,
) -> str:
    """
    Chama a API da OpenAI (Responses) para gerar o texto
    sugerido da etapa.
    """
    api_key = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
    if not api_key:
        return "⚠️ OPENAI_API_KEY não configurada."

    client = OpenAI(api_key=api_key)

    arquivos_lista = (
        ", ".join(a["nome_original"] for a in arquivos_etapa)
        if arquivos_etapa
        else "nenhum arquivo enviado"
    )

    system_prompt = (
        "Você é uma IA especialista em ETP para a Administração Pública brasileira. "
        "Gere textos claros, objetivos e alinhados à legislação de contratações públicas."
    )

    user_prompt = f"""
Informações básicas:
- Órgão: {infos_basicas.get('orgao') or '-'}
- Unidade: {infos_basicas.get('unidade') or '-'}
- Processo: {infos_basicas.get('processo') or '-'}
- Responsável: {infos_basicas.get('responsavel') or '-'}
- Objeto: {infos_basicas.get('objeto') or '-'}

Etapa: {numero_etapa} – {nome_etapa}
Orientações: {orientacao}

Arquivos de referência: {arquivos_lista}

Texto atual do usuário (se houver):
{texto_existente or '[vazio]'}

Tarefa: gere o texto final desta etapa, pronto para uso no ETP.
""".strip()

    try:
        r = client.responses.create(
            model="gpt-5",
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )

        # Tentativa 1: atributo output_text (versões mais novas)
        if hasattr(r, "output_text") and r.output_text:
            return r.output_text.strip()

        # Tentativa 2: variações de output / outputs
        outs = getattr(r, "output", None) or getattr(r, "outputs", None) or []
        partes = []
        for o in outs:
            for c in getattr(o, "content", []) or []:
                t = getattr(c, "text", None)
                if hasattr(t, "value") and t.value:
                    partes.append(t.value)
                elif isinstance(t, str):
                    partes.append(t)
                elif isinstance(t, dict):
                    partes.append(t.get("value") or t.get("text") or "")
        texto = "\n".join([p for p in partes if p]).strip()
        return texto or "⚠️ A IA não retornou texto."
    except Exception as e:
        return f"⚠️ Erro ao chamar a IA: {e}"

# ============================================================
# EXPORTAÇÃO DOCX / PDF
# ============================================================


def gerar_docx_etp(projeto: dict, etapas_rows: list) -> io.BytesIO:
    """
    Gera um DOCX em memória com todas as etapas do ETP.
    """
    doc = Document()

    doc.add_heading("Estudo Técnico Preliminar – ETP", level=0)

    doc.add_heading("Informações Básicas", level=1)
    doc.add_paragraph(f"Órgão / Entidade: {projeto.get('orgao') or ''}")
    doc.add_paragraph(f"Unidade Demandante: {projeto.get('unidade') or ''}")
    doc.add_paragraph(f"Número do Processo: {projeto.get('processo') or ''}")
    doc.add_paragraph(
        f"Responsável pela Demanda: {projeto.get('responsavel') or ''}"
    )
    doc.add_paragraph(
        f"Objeto da Contratação: {projeto.get('objeto') or ''}"
    )

    for row in etapas_rows:
        numero = row["numero"]
        titulo = row["titulo"]
        texto_final = row.get("texto_final") or "[Texto ainda não preenchido]"

        doc.add_heading(f"Etapa {numero} – {titulo}", level=1)
        for par in texto_final.split("\n\n"):
            doc.add_paragraph(par)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def gerar_pdf_etp(projeto: dict, etapas_rows: list):
    """
    Converte o DOCX gerado para PDF usando o pypandoc + wkhtmltopdf.
    Se der erro, retorna (None, mensagem_de_erro).
    """
    try:
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "etp.docx")
            pdf_path = os.path.join(tmp, "etp.pdf")

            # Salva DOCX temporário
            with open(docx_path, "wb") as f:
                f.write(gerar_docx_etp(projeto, etapas_rows).getbuffer())

            # Converte para PDF
            pypandoc.convert_file(
                docx_path,
                "pdf",
                outputfile=pdf_path,
                extra_args=["--pdf-engine=wkhtmltopdf"],
            )

            with open(pdf_path, "rb") as f:
                b = f.read()
            out = io.BytesIO(b)
            out.seek(0)
            return out, None
    except Exception as e:
        return None, str(e)

# ============================================================
# TELAS DE LOGIN / CADASTRO
# ============================================================


def tela_login_ou_cadastro():
    """
    Renderiza as abas:
    - Entrar (email/senha)
    - Cadastrar
    - Google
    """
    st.title("Ferramenta Inteligente para Elaboração de ETP")
    st.subheader("Acesse sua conta")

    tabs = st.tabs(["🔑 Entrar", "🆕 Cadastrar", "🔗 Google"])

    # --------- ENTRAR (email/senha)
    with tabs[0]:
        email = st.text_input("E-mail", key="login_email")
        senha = st.text_input(
            "Senha", type="password", key="login_senha"
        )
        col1, col2 = st.columns([1, 1])

        if col1.button("Entrar"):
            try:
                res = supabase.auth.sign_in_with_password(
                    {"email": email, "password": senha}
                )
                if res and res.session and res.session.access_token:
                    token = res.session.access_token
                    user_json = obter_user_supabase(token)
                    if user_json:
                        st.session_state["usuario"] = user_json
                        st.session_state["auth_user_id"] = user_json.get(
                            "id"
                        )
                        st.session_state["access_token"] = token
                        st.experimental_set_query_params()
                        st.rerun()
                    else:
                        st.error(
                            "Não foi possível obter os dados do usuário."
                        )
                else:
                    st.error(
                        "Falha no login. Verifique o e-mail, senha ou confirmação."
                    )
            except Exception as e:
                msg = str(e)
                if "Email not confirmed" in msg or "email not confirmed" in msg:
                    st.warning(
                        "Seu e-mail ainda não foi confirmado. Verifique sua caixa de entrada."
                    )
                else:
                    st.error(f"Erro ao autenticar: {msg}")

        

    # --------- CADASTRAR (email/senha)
    with tabs[1]:
        nome = st.text_input("Nome")
        sobrenome = st.text_input("Sobrenome")
        email_cad = st.text_input("E-mail", key="cad_email")
        senha_cad = st.text_input(
            "Senha", type="password", key="cad_senha"
        )
        if st.button("Cadastrar"):
            try:
                res = supabase.auth.sign_up(
                    {
                        "email": email_cad,
                        "password": senha_cad,
                        "options": {
                            "data": {
                                "full_name": f"{nome} {sobrenome}".strip()
                            },
                            "emailRedirectTo": APP_BASE_URL,
                        },
                    }
                )
                if res and res.user:
                    # opcionalmente, cria/atualiza o usuário local aqui:
                    existente = obter_usuario_por_email(email_cad)
                    if not existente:
                        criar_usuario(nome, sobrenome, "", email_cad)
                    st.success(
                        "Conta criada! Confirme o e-mail para conseguir entrar."
                    )
                else:
                    st.error(
                        "Não foi possível criar a conta. Verifique os dados."
                    )
            except Exception as e:
                st.error(f"Erro ao cadastrar: {e}")

    # --------- ENTRAR COM GOOGLE
    with tabs[2]:
        st.write("Ou entre com sua conta Google (via Supabase Auth).")
        auth_url = gerar_google_auth_url()
        st.link_button("🔐 Entrar com Google", auth_url)
        with st.expander("Ver URL de autenticação (debug)"):
            st.code(auth_url)

# ============================================================
# MAIN APP
# ============================================================


def main():
    # 1) Trata o callback do Google (?access_token=...)
    params = st.experimental_get_query_params()
    access_tokens = params.get("access_token")

    if "usuario" not in st.session_state and access_tokens:
        token = access_tokens[0]
        user_json = obter_user_supabase(token)
        if user_json:
            st.session_state["usuario"] = user_json
            st.session_state["auth_user_id"] = user_json.get("id")
            st.session_state["access_token"] = token
            st.experimental_set_query_params()

    # 2) Se ainda não tem usuário logado, mostra tela de login/cadastro
    if "usuario" not in st.session_state:
        tela_login_ou_cadastro()
        return

    # 3) Garante RLS ativo em toda execução
    _require_session()

    usuario = st.session_state["usuario"]

    st.title("Ferramenta Inteligente para Elaboração de ETP")

    # -------------------------------
    # Sidebar: usuário + logout
    # -------------------------------
    st.sidebar.markdown(f"**Usuário:** {usuario.get('email','')}")
    if st.sidebar.button("Sair"):
        st.session_state.clear()
        st.experimental_set_query_params()
        st.rerun()

    # -------------------------------
    # Sidebar: Projetos
    # -------------------------------
    st.sidebar.header("Projetos de ETP")
    projetos = listar_projetos()
    nomes = [p["nome"] for p in projetos]
    ids = [p["id"] for p in projetos]

    # Selectbox mostra apenas o nome (sem id)
    escolha = st.sidebar.selectbox(
        "Selecione o projeto",
        ["(Novo projeto)"] + nomes,
    )

    projeto_id = None

    if escolha == "(Novo projeto)":
        nome_novo = st.sidebar.text_input("Nome do novo projeto")
        if st.sidebar.button("Criar projeto"):
            if not nome_novo.strip():
                st.warning("Informe o nome do projeto.")
            else:
                projeto_id = criar_projeto(nome_novo.strip())
                st.success("Projeto criado com sucesso!")
                st.rerun()
    else:
        idx = nomes.index(escolha)
        projeto_id = ids[idx] if 0 <= idx < len(ids) else None

    if not projeto_id:
        st.info("Crie ou selecione um projeto de ETP na barra lateral para começar.")
        return

    projeto = obter_projeto(projeto_id)

    # -------------------------------
    # Sidebar: Gerenciar projeto (excluir)
    # -------------------------------
    st.sidebar.markdown("### Gerenciar projeto")
    confirmar = st.sidebar.checkbox(
        "Confirmar exclusão permanente", key="confirmar_exclusao"
    )
    if st.sidebar.button("🗑️ Excluir projeto selecionado"):
        if confirmar:
            excluir_projeto(projeto_id)
            st.sidebar.success("Projeto removido.")
            st.rerun()
        else:
            st.sidebar.warning(
                "Marque a caixa de confirmação antes de excluir."
            )

    # -------------------------------
    # Sidebar: Seleção de etapa
    # -------------------------------
    st.sidebar.markdown("---")
    numero_etapa = st.sidebar.selectbox(
        "Etapa",
        [n for n, _ in ETAPAS],
        format_func=lambda n: f"{n} - {dict(ETAPAS)[n]}",
    )
    nome_etapa = dict(ETAPAS)[numero_etapa]
    orientacao = ORIENTACOES.get(numero_etapa, "")

    # Status da IA
    st.sidebar.markdown("---")
    st.sidebar.caption("Status da IA:")
    if os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY"):
        st.sidebar.success("OPENAI_API_KEY configurada")
    else:
        st.sidebar.error("OPENAI_API_KEY não configurada")

    # -------------------------------
    # Layout principal
    # -------------------------------
    col1, col2 = st.columns([1.2, 2.0])

    # ---------- COLUNA ESQUERDA: INFOS + ARQUIVOS ----------
    with col1:
        st.subheader("Informações básicas do projeto")

        dados_infos = {}
        for key, label in INFOS_BASICAS_CAMPOS:
            valor_atual = (
                projeto.get(key) if projeto and projeto.get(key) is not None else ""
            )
            dados_infos[key] = st.text_input(
                label, value=valor_atual, key=f"info_{projeto_id}_{key}"
            )

        if st.button("Salvar informações básicas"):
            atualizar_infos_basicas(projeto_id, dados_infos)
            st.success("Informações atualizadas!")

        st.markdown("---")
        st.subheader(f"Arquivos da etapa {numero_etapa}")

        uploads = st.file_uploader(
            "Envie arquivos (PDF, DOCX, etc.)",
            accept_multiple_files=True,
            key=f"uploader_{projeto_id}_{numero_etapa}",
        )
        if uploads:
            for f in uploads:
                salvar_arquivo(projeto_id, numero_etapa, f)
            st.success("Arquivo(s) salvo(s).")

        lista_arquivos = listar_arquivos(projeto_id, numero_etapa)
        if lista_arquivos:
            st.caption("Arquivos cadastrados:")
            for arq in lista_arquivos:
                st.write(f"- {arq['nome_original']}")
        else:
            st.caption("Nenhum arquivo cadastrado ainda.")

    # ---------- COLUNA DIREITA: IA + TEXTO FINAL ----------
    with col2:
        st.subheader(
            f"Etapa {numero_etapa} de {len(ETAPAS)} – {nome_etapa}"
        )
        with st.expander("Orientações gerais desta etapa", expanded=True):
            st.write(orientacao)

        dados_etapa = carregar_etapa(projeto_id, numero_etapa)

        key_sug = f"sugestao_ia_{projeto_id}_{numero_etapa}"
        key_txt = f"texto_final_{projeto_id}_{numero_etapa}"

        if key_sug not in st.session_state:
            st.session_state[key_sug] = (
                dados_etapa.get("sugestao_ia", "") or ""
            )
        if key_txt not in st.session_state:
            st.session_state[key_txt] = (
                dados_etapa.get("texto_final", "") or ""
            )

        st.markdown("#### Sugestão de texto pela IA")
        if st.button(
            "Gerar sugestão com IA",
            key=f"btn_ia_{projeto_id}_{numero_etapa}",
        ):
            arquivos_etapa = [
                {"nome_original": a["nome_original"]}
                for a in listar_arquivos(projeto_id, numero_etapa)
            ]
            sugestao = gerar_texto_ia(
                numero_etapa=numero_etapa,
                nome_etapa=nome_etapa,
                orientacao=orientacao,
                texto_existente=st.session_state[key_txt],
                infos_basicas={
                    "orgao": projeto.get("orgao"),
                    "unidade": projeto.get("unidade"),
                    "processo": projeto.get("processo"),
                    "responsavel": projeto.get("responsavel"),
                    "objeto": projeto.get("objeto"),
                }
                if projeto
                else {},
                arquivos_etapa=arquivos_etapa,
            )
            st.session_state[key_sug] = sugestao

        st.text_area(
            "Sugestão da IA (edite se quiser)",
            height=200,
            key=key_sug,
        )

        st.markdown("#### Texto final da etapa")
        st.text_area(
            "Texto final que será usado no documento do ETP",
            height=300,
            key=key_txt,
        )

        if st.button(
            "Salvar etapa", key=f"btn_salvar_{projeto_id}_{numero_etapa}"
        ):
            salvar_etapa(
                projeto_id,
                numero_etapa,
                nome_etapa,
                st.session_state[key_txt],
                st.session_state[key_sug],
            )
            st.success("Etapa salva com sucesso!")

    # -------------------------------
    # Exportação DOCX / PDF
    # -------------------------------
    st.markdown("---")
    st.subheader("Exportar ETP completo")

    etapas_rows = carregar_textos_todas_etapas(projeto_id)
    if not etapas_rows:
        st.info(
            "Preencha e salve pelo menos uma etapa para habilitar a exportação."
        )
        return

    col_docx, col_pdf = st.columns(2)

    with col_docx:
        if st.button("Gerar DOCX do ETP"):
            buf = gerar_docx_etp(projeto, etapas_rows)
            st.download_button(
                "Baixar ETP em DOCX",
                data=buf,
                file_name=f"etp_projeto_{projeto_id}.docx",
                mime=(
                    "application/"
                    "vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )

    with col_pdf:
        if st.button("Gerar PDF do ETP"):
            pdf_buf, err = gerar_pdf_etp(projeto, etapas_rows)
            if err or pdf_buf is None:
                st.error(
                    "Erro ao converter DOCX para PDF no servidor. "
                    "Baixe o DOCX e converta localmente.\n" + str(err)
                )
            else:
                st.download_button(
                    "Baixar ETP em PDF",
                    data=pdf_buf,
                    file_name=f"etp_projeto_{projeto_id}.pdf",
                    mime="application/pdf",
                )


if __name__ == "__main__":
    main()
